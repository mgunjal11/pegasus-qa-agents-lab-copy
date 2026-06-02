#!/usr/bin/env python3
"""Generate Jira input template Word doc for msc-dev-code-and-qa-test-coverage-validator."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MSC-Dev-Code-and-QA-Test-Coverage-Validator-Jira-Template.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    p.add_run(text)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def build() -> Document:
    doc = Document()
    title = doc.add_heading(
        "Jira Input Template for MSC Dev Code and QA Test Coverage Validator",
        0,
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = sub.add_run("What to include on a Jira issue for accurate coverage validation reports")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Agent: ").bold = True
    meta.add_run("/msc-dev-code-and-qa-test-coverage-validator  ·  ")
    meta.add_run("Site: ").bold = True
    meta.add_run("wbdstreaming.atlassian.net  ·  ")
    meta.add_run("Author: ").bold = True
    meta.add_run("Mayur Gunjal")

    add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "The coverage validator reads Jira, linked GitHub PRs, attached Excel test plans, "
        "and optional Confluence/LADR pages. It maps acceptance criteria (R1, R2, …) to "
        "production code, dev unit/integration tests, and QA test plan cases, then produces "
        "an HTML report with coverage percentages and gap analysis."
    )
    doc.add_paragraph(
        "More complete and structured Jira content produces higher accuracy in requirement "
        "extraction, test plan mapping, and dev vs QA ownership classification."
    )

    add_heading(doc, "2. Required Jira Fields (minimum)", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Required?"
    hdr[2].text = "Why the agent needs it"
    rows = [
        ("Summary", "Yes", "Report title and story identification"),
        ("Description", "Yes", "User story scope, steps to reproduce, expected/actual behavior"),
        ("Issue Type", "Yes", "Story vs Bug affects QA scope defaults"),
        ("Status", "Yes", "Shown in report header (e.g. In QA, Ready for Release)"),
        ("Acceptance Criteria", "Strongly recommended", "Numbered R1…Rn requirements for mapping"),
        ("Attachments — test plan Excel", "Strongly recommended", "QMetry or Domino evidence sheet"),
        ("GitHub PR link(s)", "Yes for code validation", "In description or comment (full PR URL)"),
        ("Remote link — Confluence/LADR", "When applicable", "ESS milestones, partner specs"),
        ("Test data (Edit ID, Job ID, Mascot links)", "For SIT/QA AC", "Maps R4-style validation requirements"),
        ("Comments — sheet name reference", "When using SharePoint", "e.g. Refer Inc as full sheet"),
    ]
    for field, req, why in rows:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = req
        row[2].text = why

    add_heading(doc, "3. Acceptance Criteria Format (R1, R2, …)", 1)
    doc.add_paragraph(
        "Write discrete, testable acceptance criteria. The agent numbers them R1, R2, R3… "
        "automatically if you use bullets or numbered lists. Each item should be one verifiable outcome."
    )
    add_bullet(doc, " — One sentence per criterion; avoid combining unrelated checks", "Good: ")
    add_bullet(doc, " — Vague scope without measurable outcome", "Avoid: ")
    add_note(
        doc,
        "Include explicit IDs in AC when SIT validation is required, e.g. "
        "R4: Fix validated in SIT with Edit ID {uuid} and Mascot link {url}.",
    )

    add_heading(doc, "4. Description Template — User Story", 1)
    add_code_block(
        doc,
        """## User story
As a [role], I want [capability] so that [benefit].

## Acceptance criteria
R1 — [First testable outcome]
R2 — [Second testable outcome]
R3 — [Edge case or error handling]

## In scope
- [Feature A]
- [Feature B]

## Out of scope
- [Deferred item]

## Environment
- Environment: SIT / INT / PROD-like
- Platform / partner: [e.g. PFT Clear, DirecTV, ESS]

## Related design
- Confluence: [wiki URL or remote link]
- LADR: [if ESS messaging / status codes apply]

## Implementation
- PR: https://github.com/wbd-msc/{repo}/pull/{number}
- PR (additional): [repeat if multiple repos]

## Test data (for QA / SIT AC)
- Edit ID: {uuid}
- Media Request ID: {uuid}
- Mascot MVP: https://stg.foundry.wbdapps.com/mascot/fulfill-details/{id}
- Mascot incremental: [url]""",
    )

    add_heading(doc, "5. Description Template — Bug", 1)
    add_code_block(
        doc,
        """## Summary
[One-line problem statement]

## Steps to reproduce
1. [Step]
2. [Step]

## Expected behavior
[What should happen — becomes R1-style requirement]

## Actual behavior
[What happens today]

## Root cause (if known)
[Dev analysis — helps map to PR diff]

## Impact
- [Business / platform impact]

## Environment
- Environment: SIT
- Platform: [PFT Clear / TMC / CE MAM / etc.]

## Test data
- Edit ID: {uuid}
- Media Request ID: {uuid}
- Mascot links: [before/after fulfill URLs]

## Acceptance criteria (fix verification)
R1 — [Bug fixed: specific observable outcome]
R2 — [Regression: TMC/CE MAM unchanged]
R3 — [Dev: pick-genie re-fetches passport when fulfillmentType is full]
R4 — SIT validated with Edit ID {uuid} per test data above

## Implementation
- PR: https://github.com/wbd-msc/pegasus-pick-genie/pull/161""",
    )

    add_heading(doc, "6. Test Plan Attachment Requirements", 1)
    doc.add_paragraph(
        "Attach the Excel test plan on the Jira issue OR reference it in a comment with the exact sheet name."
    )

    add_heading(doc, "6.1 QMetry format", 2)
    add_bullet(doc, "Sheet name: QMetry Template")
    add_bullet(doc, "11 columns: Summary, Automatable, Automation Status, Priority, Folders, Step Summary, Test Type, Status, Regression Test (Y/N), Story, TestData Dependent")
    add_bullet(doc, "3 rows per test case with Given / When / Then in Step Summary only")
    add_bullet(doc, "Story column = Jira issue key (e.g. MSC-205625)")

    add_heading(doc, "6.2 Domino / evidence format", 2)
    add_bullet(doc, "Flexible columns: Summary, Given/When/Then or Steps, Story, Comments, SIT Jobs, QA Jobs")
    add_bullet(doc, "Include Mascot links OR Edit ID / Job ID / Pegasus ID in evidence columns")
    add_bullet(doc, "Comment on Jira: Refer [Sheet Name] sheet for Test plan and evidence")
    add_bullet(doc, "Then steps should explicitly assert the AC (e.g. passport present in cumulative output)")

    add_heading(doc, "6.3 Mapping tips for higher test plan coverage %", 2)
    add_bullet(doc, "Mention R1, R2 in test case summary or steps when possible")
    add_bullet(doc, "Use keywords from Jira AC in Then clauses (passport, fulfillmentType, frame rate, CC1, AFD, etc.)")
    add_bullet(doc, "One test case per uncovered AC — avoids R4-style gaps")
    add_bullet(doc, "For LADR/ESS stories: align scenario text to milestone names (demandAcknowledgment, orderStatus, STATUS_ERROR 9000)")

    add_heading(doc, "7. GitHub PR Linking", 1)
    doc.add_paragraph("The agent resolves PRs in this order:")
    for item in [
        "Inline --pr flag or manifest prUrls",
        "Jira description or comments — full URL: https://github.com/org/repo/pull/N",
        "Remote development links (if present)",
        "Branch search by issue key (slower; needs --repo)",
    ]:
        add_numbered(doc, item)
    add_note(
        doc,
        "Paste full PR URLs in a Jira comment labeled PR1, PR2. "
        "Partial links or branch names only reduce accuracy.",
    )

    add_heading(doc, "8. Confluence / LADR (when applicable)", 1)
    add_bullet(doc, "Add remote issue link to Confluence page OR paste wiki URL in description")
    add_bullet(doc, "For ESS/caption messaging: reference LADR ESS table (milestones + status codes 8000/9000)")
    add_bullet(doc, "For partner specs (DirecTV, etc.): link HBO Max / partner specification page")
    add_note(
        doc,
        "Without Confluence link, LADR scenarios (L1…Ln) are not extracted and test plan coverage may show gaps.",
    )

    add_heading(doc, "9. Comments Checklist (paste after QA/dev updates)", 1)
    checklist = [
        "Test plan sheet name confirmed (e.g. Inc as full → Excel tab Inc as Fulll)",
        "PR merged / open status noted",
        "SIT or INT Mascot evidence links attached or in Excel",
        "Edit ID and Job ID for SIT validation AC",
        "Known blockers linked (blocks / is blocked by issue keys)",
        "Deploy environment where fix was verified",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "10. What Improves Accuracy vs Common Gaps", 1)
    gap_table = doc.add_table(rows=1, cols=2)
    gap_table.style = "Table Grid"
    gap_table.rows[0].cells[0].text = "Common gap"
    gap_table.rows[0].cells[1].text = "Fix in Jira"
    gaps = [
        ("Test plan only on SharePoint, not attached", "Attach Excel on issue OR copy to testplans/ and note in comment"),
        ("Then steps say workflow completes but not what to verify", "Add explicit assertion matching AC wording"),
        ("SIT AC with Edit ID in description but no matching test case", "Add TC with same Edit ID in SIT Jobs column"),
        ("No PR URL on ticket", "Comment with full github.com/.../pull/N links"),
        ("Combined AC bullet with 5 checks", "Split into R1–R5 separate bullets"),
        ("LADR scenarios in test plan, no Confluence link", "Add remote link to LADR / spec page"),
        ("Multiple repos, one PR linked", "List all PR URLs in description or comment"),
    ]
    for gap, fix in gaps:
        row = gap_table.add_row().cells
        row[0].text = gap
        row[1].text = fix

    add_heading(doc, "11. Example — Well-formed Bug (MSC-205625 pattern)", 1)
    add_code_block(
        doc,
        """Summary: [PFT Clear] Content passport dropped from cumulative output manifestation (SIT)

Acceptance criteria:
R1 — Passport retained in cumulative output for PFT Clear incremental-as-full
R2 — When metadata-update becomes full fulfillment, pack passport delivered
R3 — Pick-genie must not drop passport re-fetch when fulfillmentType is full
R4 — SIT validated with Edit ID 37ea180e-77cc-413f-95cf-9dfcebf08cd2

Attachment: Domino Test Plan.xlsx
Comment: Refer Inc as full sheet for Test plan and evidence
Comment: PR1 — https://github.com/wbd-msc/pegasus-pick-genie/pull/161

Test data in description: Edit ID, Media Request ID, Mascot MVP + incremental URLs""",
    )

    add_heading(doc, "12. Running the Validator", 1)
    add_code_block(
        doc,
        """Slash command (Cursor):
  /msc-dev-code-and-qa-test-coverage-validator MSC-XXXXX

Output:
  reports/MSC-XXXXX-{date}-{time}-{TZ}.html

Re-run from cache:
  /msc-dev-code-and-qa-test-coverage-validator MSC-XXXXX --from-cache --auto""",
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = footer.add_run("MSC Dev Code and QA Test Coverage Validator · Developed by Mayur Gunjal")
    fr.italic = True
    fr.font.size = Pt(9)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
