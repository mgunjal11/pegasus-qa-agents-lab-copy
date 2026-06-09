#!/usr/bin/env python3
"""Generate Word doc: directory structure and files for msc-dev-code-and-qa-test-coverage-validator."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MSC-Dev-Code-and-QA-Test-Coverage-Validator-Directory-Guide.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)


def build() -> Document:
    doc = Document()
    title = doc.add_heading(
        "MSC Dev Code and QA Test Coverage Validator",
        0,
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = sub.add_run("Directory structure and file reference")
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.add_run("Slash command: ").bold = True
    meta.add_run("/msc-dev-code-and-qa-test-coverage-validator  ·  ")
    meta.add_run("Author: ").bold = True
    meta.add_run("Mayur Gunjal")

    doc.add_paragraph()

    add_heading(doc, "1. What this agent does", 1)
    doc.add_paragraph(
        "The agent validates an MSC Jira story against linked GitHub PR(s) (or a branch "
        "compare when no PR is linked), attached QMetry Excel test plans, and optional "
        "Confluence LADR pages. It maps acceptance criteria (R1, R2, …) to production "
        "code, dev unit/integration tests, and QA test plan cases; separates dev-owned "
        "coverage from QA handoff; and writes a timestamped HTML report with coverage "
        "percentages, traceability tables, and release readiness score."
    )
    add_bullet(doc, " — end-to-end, no mid-run confirmation", "Default invoke")
    add_code(doc, "/msc-dev-code-and-qa-test-coverage-validator MSC-205625")

    add_heading(doc, "2. High-level directory tree", 1)
    add_code(
        doc,
        """<workspace-root>/
├── AGENTS.md                          # Workspace pointer to this agent + skill
├── .env.example                       # Jira credentials template (optional fetches)
├── .coverage-validator.defaults.json  # Optional repo defaults (not committed)
├── .cursor/
│   ├── agents/
│   │   └── msc-dev-code-and-qa-test-coverage-validator.md
│   ├── commands/
│   │   └── msc-dev-code-and-qa-test-coverage-validator.md
│   ├── permissions.json               # MCP + terminal allowlist (after install)
│   └── skills/
│       └── msc-dev-code-and-qa-test-coverage-validator/
│           ├── SKILL.md               # Primary workflow (Steps 0–9)
│           ├── report-template.html   # HTML shell with {{PLACEHOLDER}} tokens
│           ├── report-template.md     # Placeholder documentation
│           ├── examples.md
│           ├── validator.defaults.example.json
│           └── references/
│               ├── auto-approve-setup.md
│               ├── run-options.md
│               ├── dev-qa-test-scope.md
│               ├── jira-testplan-validation.md
│               ├── confluence-ladr-requirements.md
│               └── github-coverage.md
├── scripts/                           # Python pipeline (see Section 4)
├── docs/
│   ├── MSC-Dev-Code-and-QA-Test-Coverage-Validator-Jira-Template.docx
│   └── MSC-Dev-Code-and-QA-Test-Coverage-Validator-Directory-Guide.docx  (this file)
└── reports/
    ├── <KEY>-<date>-<TZ>.html         # Output reports
    └── .cache/
        ├── <KEY>-manifest.json
        ├── <KEY>-jira.json
        ├── <KEY>-confluence.json
        ├── <KEY>-testplan.json
        ├── <KEY>-prefetch.json
        ├── <KEY>-mapping.json
        ├── <KEY>-analysis.json      # Optional manual overrides
        └── <KEY>-testplan-files/      # Downloaded Excel attachments""",
    )

    add_heading(doc, "3. Cursor agent layer (how Cursor invokes it)", 1)
    add_table(
        doc,
        ["File", "Role"],
        [
            [
                ".cursor/commands/msc-dev-code-and-qa-test-coverage-validator.md",
                "Slash command definition: pipeline steps 0–11, --auto --write behavior.",
            ],
            [
                ".cursor/agents/msc-dev-code-and-qa-test-coverage-validator.md",
                "Subagent definition: auto-run rules, report builder notes, hard rules.",
            ],
            [
                ".cursor/skills/.../SKILL.md",
                "Full skill: metrics, workflow, MCP usage, cache paths, HTML placeholders.",
            ],
            [
                "AGENTS.md",
                "Workspace summary; links skill path and report output pattern.",
            ],
            [
                ".cursor/permissions.json",
                "Created by install_coverage_validator_permissions.py — reduces Allow prompts.",
            ],
        ],
    )

    add_heading(doc, "4. Python scripts (pipeline order)", 1)
    doc.add_paragraph(
        "Run from workspace root. The agent typically executes these in order for each Jira key:"
    )
    add_table(
        doc,
        ["Step", "Script", "Purpose"],
        [
            [
                "0",
                "(manifest / defaults)",
                "Merge flags → reports/.cache/{KEY}-manifest.json and .coverage-validator.defaults.json",
            ],
            [
                "2",
                "Atlassian MCP (agent)",
                "getJiraIssue + getJiraIssueRemoteIssueLinks → persist {KEY}-jira.json",
            ],
            [
                "3",
                "fetch_confluence_requirements.py",
                "LADR/Confluence requirements → {KEY}-confluence.json",
            ],
            [
                "4",
                "fetch_jira_testplan.py",
                "Download/parse Excel test plan → {KEY}-testplan.json",
            ],
            [
                "5",
                "prefetch_coverage_inputs.py or fetch_coverage_github.py",
                "GitHub PR diff, CI, or branch compare → {KEY}-prefetch.json",
            ],
            [
                "6",
                "map_requirements_to_diff.py",
                "Score R-items vs code/tests → {KEY}-mapping.json",
            ],
            [
                "7",
                "build_coverage_report.py",
                "Fill report-template.html, apply_report_ui_enhancements() → HTML report",
            ],
        ],
    )

    add_heading(doc, "4.1 Core library modules", 1)
    add_table(
        doc,
        ["File", "Used by", "Purpose"],
        [
            [
                "coverage_report_helpers.py",
                "build_coverage_report.py",
                "Report HTML builders, tooltips (v22), Jira readiness UI, PR rows, §3/§4 fields, apply_report_ui_enhancements()",
            ],
            [
                "coverage_report_timestamp.py",
                "build_coverage_report.py",
                "Timestamped report path: reports/{KEY}-{MM-DD-YYYY-HH-MM-SS}-{TZ}.html",
            ],
            [
                "ci_coverage.py",
                "prefetch_coverage_inputs.py, build_coverage_report.py",
                "Codecov/Sonar/pytest-cov; Sonar estimated-after-merge PR comments; HTTP 410 / expired artifact fallback",
            ],
            [
                "confluence_requirements.py",
                "fetch_confluence_requirements.py",
                "Parse LADR requirements; dedupe_ladr_requirements(); unique-id test plan coverage %",
            ],
            [
                "testplan_gwt.py",
                "fetch_jira_testplan.py",
                "Given/When/Then parsing from Excel test plan",
            ],
            [
                "testplan_evidence.py",
                "fetch_jira_testplan.py, coverage_report_helpers.py",
                "Mascot links, SIT job IDs in test plan Evidence column",
            ],
            [
                "jira_env.py",
                "fetch_jira_testplan.py, verify_jira_credentials.py",
                "Jira API credentials from .env",
            ],
        ],
    )

    add_heading(doc, "4.2 Supporting / optional scripts", 1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["install_coverage_validator_permissions.py", "One-time Cursor allowlist install"],
            ["generate_jira_template_for_coverage_validator.py", "Jira input quality Word template"],
            ["generate_coverage_validator_directory_guide.py", "Generates this directory guide Word doc"],
            ["generate_coverage_validator_ppt.py", "Optional management PPT deck from report HTML"],
            ["ppt_report_from_html.py", "Helper for PPT generation"],
            ["sync_pegasus_qa_agents_lab.py", "Publish agent + scripts to pegasus-qa-agents-lab GitHub repo"],
            ["verify_jira_credentials.py", "Test Jira .env configuration"],
            ["upload_jira_testplan.py", "Optional: upload test plan to Jira"],
            ["regen_msc205625_report.py, regen_msc204417_report.py", "Legacy per-ticket regen scripts (prefer build_coverage_report.py)"],
        ],
    )

    add_heading(doc, "4.3 Test scripts (regression)", 1)
    doc.add_paragraph("Run before releases or after UI/helper changes:")
    add_code(
        doc,
        "python -m pytest scripts/test_report_ui_enhancements.py scripts/test_quick_links.py "
        "scripts/test_qa_scope_handoff.py scripts/test_confluence_requirements.py -q",
    )
    add_bullet(doc, " — tooltip layout v22, idempotent UI injectors", "test_report_ui_enhancements.py")
    add_bullet(doc, " — LADR quick links filtering", "test_quick_links.py")
    add_bullet(doc, " — qaScope: none when dev tests covered", "test_qa_scope_handoff.py")
    add_bullet(doc, " — Confluence LADR parsing", "test_confluence_requirements.py")
    add_bullet(doc, " — PR table, CI fields, mapping, GWT, evidence", "Also: test_pr_rows, test_ci_*, test_map_requirements_to_diff, test_gwt_parsing, test_testplan_evidence")

    add_heading(doc, "5. Skill folder files (templates & references)", 1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["report-template.html", "HTML report skeleton; build_coverage_report.py replaces {{PLACEHOLDER}} tokens"],
            ["report-template.md", "Documents all placeholders ({{PR_ROWS}}, {{REQUIREMENT_ROWS}}, etc.)"],
            ["validator.defaults.example.json", "Example .coverage-validator.defaults.json"],
            ["references/run-options.md", "--auto, --from-cache, --fetch-only, manifest flags"],
            ["references/auto-approve-setup.md", "Cursor permissions / one-shot pipeline setup"],
            ["references/dev-qa-test-scope.md", "Dev vs QA ownership rules"],
            ["references/jira-testplan-validation.md", "Excel test plan expectations"],
            ["references/confluence-ladr-requirements.md", "LADR L1…Ln traceability"],
            ["references/github-coverage.md", "PR prefetch, branch compare, gh usage"],
            ["examples.md", "Sample invocations and output paths"],
        ],
    )

    add_heading(doc, "6. Runtime outputs and cache files", 1)
    add_table(
        doc,
        ["Path", "Written by", "Contents"],
        [
            ["reports/{KEY}-*.html", "build_coverage_report.py", "Final HTML report (UI v22 applied)"],
            ["reports/.cache/{KEY}-manifest.json", "prefetch / builder", "lastReportFile, prUrls, repo, mode, timezone"],
            ["reports/.cache/{KEY}-jira.json", "Agent MCP + scripts", "Requirements R1…Rn, attachments, PR URLs, Confluence links"],
            ["reports/.cache/{KEY}-confluence.json", "fetch_confluence_requirements.py", "LADR requirements L1…Ln"],
            ["reports/.cache/{KEY}-testplan.json", "fetch_jira_testplan.py", "Parsed test cases, GWT, Jira/LADR mapping"],
            ["reports/.cache/{KEY}-prefetch.json", "prefetch / fetch_coverage_github", "PR diffs, CI status, branchCompare files"],
            ["reports/.cache/{KEY}-mapping.json", "map_requirements_to_diff.py", "Per-R code/dev-test scores, qaScope, evidence files"],
            ["reports/.cache/{KEY}-analysis.json", "Manual (optional)", "Override narrative lists and coverage % for build_coverage_report.py --analysis"],
        ],
    )

    add_heading(doc, "7. External tools (not in repo)", 1)
    add_bullet(doc, " — getJiraIssue, getJiraIssueRemoteIssueLinks, getConfluencePage", "Atlassian MCP")
    add_bullet(doc, " — PR diff, checks, branch compare (via prefetch scripts)", "GitHub CLI (gh)")
    add_bullet(doc, " — Jira attachment download when not using MCP-only cache", "Jira REST (.env)")

    add_heading(doc, "8. End-to-end pipeline (one page)", 1)
    add_code(
        doc,
        """/msc-dev-code-and-qa-test-coverage-validator MSC-205625

  MCP (parallel)     →  reports/.cache/MSC-205625-jira.json
  fetch_confluence   →  reports/.cache/MSC-205625-confluence.json
  fetch_jira_testplan→  reports/.cache/MSC-205625-testplan.json
  prefetch / gh      →  reports/.cache/MSC-205625-prefetch.json
  map_requirements   →  reports/.cache/MSC-205625-mapping.json
  build_coverage_report
                     →  reports/MSC-205625-<timestamp>-IST.html
                     →  manifest lastReportFile updated""",
    )

    add_heading(doc, "9. Published copy (pegasus-qa-agents-lab)", 1)
    doc.add_paragraph(
        "A mirror of this agent ships to GitHub: github.com/mgunjal11/pegasus-qa-agents-lab. "
        "Run scripts/sync_pegasus_qa_agents_lab.py to copy .cursor agents/skills, scripts, "
        "and AGENTS.md into that repo for teammates."
    )

    add_heading(doc, "10. Report UI notes (current)", 1)
    add_bullet(doc, " — hover i icons; layout v22", "Tooltips")
    add_bullet(doc, " — no tooltip on report h1 or Jira input readiness h3", "Excluded tooltips")
    add_bullet(doc, " — checklist rows (AC, PR, test plan, LADR) still have tooltips", "Jira readiness")
    add_bullet(doc, " — green check / red X on readiness rows", "Icons")

    note = doc.add_paragraph()
    note.add_run("Regenerate this document: ").bold = True
    note.add_run("python scripts/generate_coverage_validator_directory_guide.py")

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
